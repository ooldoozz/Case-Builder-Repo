from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps

from .common import image_file_path
from .preview_defaults import width_span

STATUS_COLORS = {
    "complete": "15803D",
    "weak": "C2410C",
    "unclear": "6D28D9",
    "missing": "B91C1C",
}


def _borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _shading(cell, fill: str | None) -> None:
    if not fill:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def _margins(cell, value: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _image_source(
    case_id: int,
    image: dict,
    width_inches: float,
    height_inches: float,
    fit: str,
) -> BytesIO | None:
    path = image_file_path(case_id, image.get("location"))
    if path is None:
        return None

    try:
        if path.suffix.lower() == ".svg":
            try:
                import cairosvg
            except ImportError:
                return None
            raw = BytesIO(cairosvg.svg2png(url=str(path)))
        else:
            raw = BytesIO(path.read_bytes())

        with Image.open(raw) as source:
            source = source.convert("RGBA")
            target_width = max(240, int(width_inches * 144))
            target_height = max(140, int(height_inches * 144))

            if fit == "contain":
                contained = source.copy()
                contained.thumbnail((target_width, target_height), Image.Resampling.LANCZOS)
                canvas = Image.new("RGBA", (target_width, target_height), (248, 250, 252, 255))
                left = (target_width - contained.width) // 2
                top = (target_height - contained.height) // 2
                canvas.alpha_composite(contained, (left, top))
                prepared = canvas
            else:
                prepared = ImageOps.fit(
                    source,
                    (target_width, target_height),
                    method=Image.Resampling.LANCZOS,
                    centering=(0.5, 0.5),
                )

            output = BytesIO()
            prepared.convert("RGB").save(output, format="PNG", optimize=True)
            output.seek(0)
            return output
    except Exception:
        return None


def _add_picture(
    paragraph,
    case_id: int,
    image: dict,
    width_inches: float,
    height_inches: float,
    fit: str,
) -> bool:
    source = _image_source(case_id, image, width_inches, height_inches, fit)
    if source is None:
        return False

    try:
        paragraph.add_run().add_picture(
            source,
            width=Inches(max(1.0, width_inches)),
            height=Inches(max(0.8, height_inches)),
        )
        return True
    except Exception:
        return False


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _add_media(cell, case_id: int, images: list[dict], settings: dict, width_inches: float) -> None:
    if not images:
        return

    columns = max(1, min(2, int(settings.get("image_columns", 1))))
    columns = min(columns, len(images))
    grid = cell.add_table(rows=0, cols=columns)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False

    gap_allowance = 0.12 * max(0, columns - 1)
    item_width = max(1.0, (width_inches - gap_allowance) / columns)
    image_height = max(0.8, float(settings.get("image_height", 230)) / 96.0)
    image_fit = settings.get("image_fit", "cover")

    for offset in range(0, len(images), columns):
        row = grid.add_row()
        _set_row_cant_split(row)
        row_images = images[offset:offset + columns]

        for index, media_cell in enumerate(row.cells):
            _borderless(media_cell)
            _margins(media_cell, 35)
            media_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            media_cell.width = Inches(item_width)

            if index >= len(row_images):
                continue

            image = row_images[index]
            picture = media_cell.paragraphs[0]
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not _add_picture(
                picture,
                case_id,
                image,
                item_width,
                image_height,
                image_fit,
            ):
                continue

            if settings.get("show_evidence_type") and image.get("evidence_type"):
                label = media_cell.add_paragraph()
                label.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label.paragraph_format.space_before = Pt(3)
                label.paragraph_format.space_after = Pt(1)
                run = label.add_run(str(image["evidence_type"]))
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(102, 112, 133)

            if settings.get("show_caption") and image.get("caption"):
                caption = media_cell.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(5)
                run = caption.add_run(str(image["caption"]))
                run.italic = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(102, 112, 133)

def _add_text(cell, section: dict, settings: dict, show_status: bool) -> None:
    heading = cell.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.keep_with_next = True

    number = heading.add_run(f'{section["number"]}  ')
    number.font.size = Pt(8)
    number.font.color.rgb = RGBColor(152, 162, 179)

    title = heading.add_run(section["title"])
    title.bold = True
    title.font.size = Pt(settings["title_size"])
    title.font.color.rgb = RGBColor(16, 24, 40)

    if show_status:
        status = heading.add_run(f'   • {section["status"].title()}')
        status.bold = True
        status.font.size = Pt(8)
        status.font.color.rgb = RGBColor.from_string(
            STATUS_COLORS.get(section["status"], "667085")
        )

    body = section.get("body") or ""
    if not body:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run("No content available for this section.")
        run.italic = True
        run.font.size = Pt(settings["body_size"])
        run.font.color.rgb = RGBColor(185, 28, 28)
        return

    for text in body.split("\n"):
        if not text.strip():
            continue
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = settings["line_height"]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if settings.get("text_align") == "center"
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        run = paragraph.add_run(text.strip())
        run.font.size = Pt(settings["body_size"])
        run.font.color.rgb = RGBColor(71, 84, 103)


def _add_section_content(
    cell,
    case: Any,
    section: dict,
    images: list[dict],
    settings: dict,
    show_status: bool,
    width_inches: float,
) -> None:
    layout = settings.get("layout", "text_only")

    if layout in {"media_left", "media_right"} and images:
        split = cell.add_table(rows=1, cols=2)
        split.autofit = False
        for split_cell in split.rows[0].cells:
            _borderless(split_cell)
            _margins(split_cell, 40)
            split_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        text_cell, media_cell = split.rows[0].cells
        if layout == "media_left":
            text_cell, media_cell = media_cell, text_cell
        _add_text(text_cell, section, settings, show_status)
        _add_media(media_cell, case.id, images, settings, max(1.1, width_inches / 2 - 0.2))
        return

    if layout == "media_top" and images:
        _add_media(cell, case.id, images, settings, max(1.2, width_inches - 0.3))
    _add_text(cell, section, settings, show_status)
    if layout == "media_bottom" and images:
        _add_media(cell, case.id, images, settings, max(1.2, width_inches - 0.3))


def _group_rows(sections: list[dict], preview: dict) -> list[list[dict]]:
    rows: list[list[dict]] = []
    current: list[dict] = []
    used = 0

    for section in sections:
        if (
            section["status"] == "missing"
            and not preview["document"].get("include_missing_sections", True)
        ):
            continue
        span = width_span(preview["sections"][section["key"]]["width"])
        if current and used + span > 12:
            rows.append(current)
            current = []
            used = 0
        current.append(section)
        used += span
        if used >= 12:
            rows.append(current)
            current = []
            used = 0

    if current:
        rows.append(current)
    return rows


def build_case_docx(
    case: Any,
    sections: list[dict[str, Any]],
    all_images: list[dict[str, Any]],
    preview: dict[str, Any],
) -> bytes:
    document = Document()
    page = document.sections[0]
    page.top_margin = Inches(0.55)
    page.bottom_margin = Inches(0.55)
    page.left_margin = Inches(0.58)
    page.right_margin = Inches(0.58)

    font_family = preview["document"]["font_family"]
    normal = document.styles["Normal"]
    normal.font.name = font_family
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_family)
    normal.font.size = Pt(preview["document"]["base_font_size"])

    eyebrow = document.add_paragraph()
    run = eyebrow.add_run(preview["cover"]["eyebrow"])
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(102, 112, 133)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(case.title or "Untitled Case Study")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(16, 24, 40)

    if preview["cover"].get("subtitle"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(14)
        run = paragraph.add_run(preview["cover"]["subtitle"])
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(71, 84, 103)

    meta = document.add_table(rows=1, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta_items = [
        ("ROLE", preview["cover"].get("role", "")),
        ("PLATFORM", preview["cover"].get("platform", "")),
        ("FOCUS", preview["cover"].get("focus", "")),
        ("TIMELINE", preview["cover"].get("timeline", "")),
    ]
    for index, (label, value) in enumerate(meta_items):
        cell = meta.rows[0].cells[index]
        _borderless(cell)
        _shading(cell, "F8FAFC")
        _margins(cell, 120)
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.size = Pt(7)
        r1.font.color.rgb = RGBColor(152, 162, 179)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(value or "—")
        r2.bold = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(52, 64, 84)

    images_by_id = {image["id"]: image for image in all_images}
    cover_image = images_by_id.get(preview["cover"].get("cover_image_id"))
    if cover_image:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_picture(
            paragraph,
            case.id,
            cover_image,
            6.85,
            max(1.65, float(preview["cover"].get("image_height", 300)) / 96.0),
            preview["cover"].get("image_fit", "cover"),
        )

    document.add_paragraph()

    for row_sections in _group_rows(sections, preview):
        table = document.add_table(rows=1, cols=len(row_sections))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_row_cant_split(table.rows[0])

        for index, section in enumerate(row_sections):
            settings = preview["sections"][section["key"]]
            span = width_span(settings["width"])
            width_inches = 7.0 * span / 12
            cell = table.rows[0].cells[index]
            cell.width = Inches(width_inches)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _borderless(cell)
            _margins(cell, settings["padding"] * 10)
            if settings["style"] == "soft":
                _shading(cell, "F8FAFC")

            images = [
                images_by_id[image_id]
                for image_id in settings.get("image_order", [])
                if image_id in images_by_id
                and images_by_id[image_id].get("section_key") == section["key"]
            ]
            _add_section_content(
                cell,
                case,
                section,
                images,
                settings,
                preview["document"].get("show_status_badges", False),
                width_inches,
            )
        document.add_paragraph().paragraph_format.space_after = Pt(3)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
